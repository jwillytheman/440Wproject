from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
def stripline(line):
    newline = ''.join(line.rsplit("<", 2))
    evennewerline = ' '.join(newline.rsplit("/", 2))
    newerline = ' '.join(evennewerline.rsplit(">", 2))
    return newerline

def paragraphCreation():
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph_format.space_before = Pt(18)
    return paragraph

document = Document()
inputText = raw_input('Please input the text of the first question. \n '
                      'When all questions are put in, type "quit"')
while inputText != 'quit':
    paragraph1 = paragraphCreation()
    paragraph1.add_run(inputText)
    inputText = raw_input('Please input the text of the next question. \n '
                          'When all questions are put in, type "quit"\n')

document.save('test.docx')
with open("TALA_Test_Eval.xml","r") as file:
    check = 0
    checkText = []
    for line in file:
        if '<Q' in line:
            checkText.append(stripline(line))
        if 'Response' in line and check == 0:
            check = 1
        elif 'Response' in line and check == 1:
            check = 0
        elif check == 1:
            for entry in checkText:
                if entry in line:
                    newline = stripline(line)
                    newline = ''.join(newline.rsplit(entry))
                    newline = ''.join(newline.strip('	'))
                    print "Response to " + entry + " is " + newline







